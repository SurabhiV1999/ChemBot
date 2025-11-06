"""
Document Processing Module
Converts various file formats (PDF, DOCX, TXT, MD) to Markdown format
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any
import pymupdf4llm
from docx import Document as DocxDocument
from pdf2docx import Converter
import markdown

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process and convert documents to Markdown format"""
    
    def __init__(self):
        self.temp_dir = "temp_conversions"
        os.makedirs(self.temp_dir, exist_ok=True)
    
    async def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Process document and convert to Markdown
        
        Args:
            file_path: Path to the document
            file_type: MIME type of the document
        
        Returns:
            Dict with markdown_text, metadata, and processing info
        """
        try:
            logger.info(f"Processing document: {file_path} (type: {file_type})")
            
            if file_type == "application/pdf":
                return await self._process_pdf(file_path)
            elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                return await self._process_docx(file_path)
            elif file_type == "text/plain":
                return await self._process_txt(file_path)
            elif file_type == "text/markdown":
                return await self._process_md(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Convert PDF to Markdown using PyMuPDF4LLM"""
        try:
            logger.info("Converting PDF to Markdown using PyMuPDF4LLM")
            
            # Use PyMuPDF4LLM to extract markdown
            markdown_text = pymupdf4llm.to_markdown(file_path)
            
            # Get metadata
            metadata = self._extract_pdf_metadata(file_path)
            
            return {
                "markdown_text": markdown_text,
                "metadata": metadata,
                "page_count": metadata.get("page_count", 0),
                "word_count": len(markdown_text.split()),
                "conversion_method": "pymupdf4llm"
            }
        
        except Exception as e:
            logger.error(f"Error converting PDF: {str(e)}")
            raise
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Convert DOCX to Markdown (via PDF conversion)"""
        try:
            logger.info("Converting DOCX to Markdown")
            
            # First convert DOCX to PDF
            pdf_path = os.path.join(self.temp_dir, f"{Path(file_path).stem}.pdf")
            
            # Convert DOCX to PDF
            cv = Converter(file_path)
            cv.convert(pdf_path)
            cv.close()
            
            # Now convert PDF to Markdown
            result = await self._process_pdf(pdf_path)
            result["conversion_method"] = "docx->pdf->markdown"
            
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            
            return result
        
        except Exception as e:
            logger.error(f"Error converting DOCX: {str(e)}")
            # Fallback: extract text directly from DOCX
            return await self._process_docx_fallback(file_path)
    
    async def _process_docx_fallback(self, file_path: str) -> Dict[str, Any]:
        """Fallback method to extract text from DOCX"""
        try:
            logger.info("Using fallback method for DOCX")
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Try to detect headings and format accordingly
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        try:
                            level_num = int(level)
                            paragraphs.append(f"{'#' * level_num} {para.text}")
                        except ValueError:
                            paragraphs.append(f"## {para.text}")
                    else:
                        paragraphs.append(para.text)
            
            markdown_text = "\n\n".join(paragraphs)
            
            return {
                "markdown_text": markdown_text,
                "metadata": {"source": "docx", "paragraphs": len(doc.paragraphs)},
                "page_count": 1,  # DOCX doesn't have pages
                "word_count": len(markdown_text.split()),
                "conversion_method": "docx_direct"
            }
        
        except Exception as e:
            logger.error(f"Error in DOCX fallback: {str(e)}")
            raise
    
    async def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """Convert TXT to Markdown"""
        try:
            logger.info("Converting TXT to Markdown")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Simple conversion: add markdown formatting
            lines = text.split('\n')
            markdown_lines = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    markdown_lines.append("")
                    continue
                
                # Try to detect headings (lines that are short and followed by content)
                if len(line) < 100 and line.isupper():
                    markdown_lines.append(f"## {line}")
                elif len(line) < 100 and not line.endswith(('.', ',', ';', ':')):
                    markdown_lines.append(f"### {line}")
                else:
                    markdown_lines.append(line)
            
            markdown_text = "\n\n".join(markdown_lines)
            
            return {
                "markdown_text": markdown_text,
                "metadata": {"source": "txt", "lines": len(lines)},
                "page_count": 1,
                "word_count": len(markdown_text.split()),
                "conversion_method": "txt_direct"
            }
        
        except Exception as e:
            logger.error(f"Error converting TXT: {str(e)}")
            raise
    
    async def _process_md(self, file_path: str) -> Dict[str, Any]:
        """Process Markdown file (already in correct format)"""
        try:
            logger.info("Reading Markdown file")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                markdown_text = f.read()
            
            return {
                "markdown_text": markdown_text,
                "metadata": {"source": "markdown"},
                "page_count": 1,
                "word_count": len(markdown_text.split()),
                "conversion_method": "direct"
            }
        
        except Exception as e:
            logger.error(f"Error reading Markdown: {str(e)}")
            raise
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF"""
        try:
            import pymupdf
            doc = pymupdf.open(file_path)
            
            metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "keywords": doc.metadata.get("keywords", ""),
            }
            
            doc.close()
            return metadata
        
        except Exception as e:
            logger.warning(f"Could not extract PDF metadata: {str(e)}")
            return {}
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception as e:
            logger.warning(f"Error cleaning up temp directory: {str(e)}")

